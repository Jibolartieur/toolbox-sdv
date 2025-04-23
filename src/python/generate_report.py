import sys
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Charger les données JSON depuis l'entrée standard
input_data = json.load(sys.stdin)

if 'results' not in input_data or 'target' not in input_data:
    print('Error: No results or target provided')
    sys.exit(1)

# Créer un nouveau document Word
doc = Document()

# Ajouter une page de garde
doc.add_heading('Synthèse de l\'audit de sécurité', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph(f"Cible : {input_data['target']}", style='Title').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph(f"Date : {input_data.get('date', 'Non spécifiée')}", style='Title').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_page_break()

# Ajouter les résultats des outils
results = input_data['results'].split("\n\n")

for result in results:
    if '=== ' in result:
        tool_name = result.split('=== ')[1].split(' Results ===')[0]
        doc.add_heading(f"{tool_name} Results", level=2)
        content = result.split('=== ')[1].split(' Results ===')[1].strip()
        doc.add_paragraph(content, style='BodyText')

# Enregistrer le document
output_filename = 'security_audit_report.docx'
doc.save(output_filename)

print(f'Report generated: {output_filename}') 